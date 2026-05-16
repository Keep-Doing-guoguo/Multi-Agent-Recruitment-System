from __future__ import annotations

from pathlib import Path
from typing import Protocol
from urllib.parse import urlparse

from recruitment_system.models import DocumentExtractionResult, DocumentPurpose


class MultimodalExtractor(Protocol):
    """Adapter interface for model-backed document understanding."""

    def extract(self, source: str | Path, purpose: DocumentPurpose) -> DocumentExtractionResult:
        ...


class DocumentExtractionAgent:
    """Converts raw file/text input into standardized text for downstream agents."""

    TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".csv", ".json"}
    PDF_EXTENSIONS = {".pdf"}
    DOCX_EXTENSIONS = {".docx"}

    def __init__(self, multimodal_extractor: MultimodalExtractor | None = None) -> None:
        self.multimodal_extractor = multimodal_extractor

    def run(self, source: str, purpose: DocumentPurpose) -> DocumentExtractionResult:
        if not source.strip():
            return DocumentExtractionResult(
                source=source,
                purpose=purpose,
                confidence=0.0,
                errors=[f"{purpose}_input is required"],
            )

        if self._is_url(source):
            if self.multimodal_extractor is not None:
                return self.multimodal_extractor.extract(source, purpose)
            return DocumentExtractionResult(
                source=source,
                purpose=purpose,
                file_type=self._file_type_from_source(source),
                confidence=0.0,
                warnings=["需要配置多模态提取器后才能解析 URL 文件"],
                errors=["multimodal_extractor_required"],
            )

        path = Path(source)
        if path.exists() and path.is_file():
            return self._extract_file(path, purpose)

        return DocumentExtractionResult(
            source="inline_text",
            purpose=purpose,
            file_type="text",
            extracted_text=source,
            confidence=1.0,
        )

    def _extract_file(self, path: Path, purpose: DocumentPurpose) -> DocumentExtractionResult:
        suffix = path.suffix.lower()
        if suffix in self.TEXT_EXTENSIONS:
            return DocumentExtractionResult(
                source=str(path),
                purpose=purpose,
                file_type=suffix.lstrip(".") or "text",
                extracted_text=path.read_text(encoding="utf-8"),
                confidence=1.0,
            )
        if suffix in self.PDF_EXTENSIONS:
            return self._extract_pdf(path, purpose)
        if suffix in self.DOCX_EXTENSIONS:
            return self._extract_docx(path, purpose)

        if self.multimodal_extractor is not None:
            return self.multimodal_extractor.extract(path, purpose)

        return DocumentExtractionResult(
            source=str(path),
            purpose=purpose,
            file_type=suffix.lstrip(".") or "unknown",
            confidence=0.0,
            warnings=["需要配置多模态提取器后才能解析该文件类型"],
            errors=[f"unsupported_file_type: {suffix or 'unknown'}"],
        )

    def _is_url(self, source: str) -> bool:
        parsed = urlparse(source)
        return parsed.scheme in {"http", "https"} and bool(parsed.netloc)

    def _file_type_from_source(self, source: str) -> str:
        suffix = Path(urlparse(source).path).suffix.lower()
        return suffix.lstrip(".") or "url"

    def _extract_pdf(self, path: Path, purpose: DocumentPurpose) -> DocumentExtractionResult:
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            pages = [page.extract_text() or "" for page in reader.pages]
            extracted_text = "\n\n".join(page.strip() for page in pages if page.strip())
            if not extracted_text and self.multimodal_extractor is not None:
                return self.multimodal_extractor.extract(path, purpose)
            warnings = [] if extracted_text else ["PDF 未提取到文本，可能是扫描件，需要 OCR 或多模态解析"]
            return DocumentExtractionResult(
                source=str(path),
                purpose=purpose,
                file_type="pdf",
                extracted_text=extracted_text,
                confidence=0.95 if extracted_text else 0.0,
                layout_blocks=[f"page_{index + 1}" for index, page in enumerate(pages) if page.strip()],
                warnings=warnings,
            )
        except Exception as error:
            if self.multimodal_extractor is not None:
                result = self.multimodal_extractor.extract(path, purpose)
                result.warnings.append(f"本地 PDF 解析失败，已回退到多模态解析：{error}")
                return result
            return DocumentExtractionResult(
                source=str(path),
                purpose=purpose,
                file_type="pdf",
                confidence=0.0,
                errors=[f"pdf_extraction_failed: {error}"],
            )

    def _extract_docx(self, path: Path, purpose: DocumentPurpose) -> DocumentExtractionResult:
        try:
            from docx import Document

            document = Document(str(path))
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            table_rows: list[str] = []
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        table_rows.append(" | ".join(cells))
            extracted_text = "\n".join(paragraphs + table_rows)
            warnings = [] if extracted_text else ["DOCX 未提取到文本"]
            return DocumentExtractionResult(
                source=str(path),
                purpose=purpose,
                file_type="docx",
                extracted_text=extracted_text,
                confidence=0.95 if extracted_text else 0.0,
                tables=table_rows,
                warnings=warnings,
            )
        except Exception as error:
            return DocumentExtractionResult(
                source=str(path),
                purpose=purpose,
                file_type="docx",
                confidence=0.0,
                errors=[f"docx_extraction_failed: {error}"],
            )
