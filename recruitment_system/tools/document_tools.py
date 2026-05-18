from __future__ import annotations

from pathlib import Path

from recruitment_system.models import DocumentExtractionResult, DocumentPurpose


class TextParserTool:
    """读取文本类文件。"""

    def parse(self, path: Path, purpose: DocumentPurpose) -> DocumentExtractionResult:
        """按 UTF-8 读取文本文件并封装成 DocumentExtractionResult。"""
        return DocumentExtractionResult(
            source=str(path),
            purpose=purpose,
            file_type=path.suffix.lower().lstrip(".") or "text",
            extracted_text=path.read_text(encoding="utf-8"),
            confidence=1.0,
        )


class PdfParserTool:
    """从带文本层的 PDF 中提取文本。"""

    def parse(self, path: Path, purpose: DocumentPurpose) -> DocumentExtractionResult:
        """使用 pypdf 提取 PDF 文本；扫描件会返回空文本警告。"""
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            pages = [page.extract_text() or "" for page in reader.pages]
            extracted_text = "\n\n".join(page.strip() for page in pages if page.strip())
            warnings = [] if extracted_text else ["PDF 未提取到文本，可能是扫描件，需要 OCR 或多模态解析"]
            return DocumentExtractionResult(
                source=str(path),
                purpose=purpose,
                file_type="pdf",
                extracted_text=extracted_text,
                confidence=0.95 if extracted_text else 0.0,
                layout_blocks=[f"page_{index + 1}" for index, page in enumerate(pages) if page.strip()],
                warnings=warnings,
            )
        except Exception as error:
            return DocumentExtractionResult(
                source=str(path),
                purpose=purpose,
                file_type="pdf",
                confidence=0.0,
                errors=[f"pdf_extraction_failed: {error}"],
            )


class DocxParserTool:
    """从 DOCX 段落和表格中提取文本。"""

    def parse(self, path: Path, purpose: DocumentPurpose) -> DocumentExtractionResult:
        """读取 DOCX 的段落和表格内容并合并为纯文本。"""
        try:
            from docx import Document

            document = Document(str(path))
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            table_rows: list[str] = []
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        table_rows.append(" | ".join(cells))
            extracted_text = "\n".join(paragraphs + table_rows)
            warnings = [] if extracted_text else ["DOCX 未提取到文本"]
            return DocumentExtractionResult(
                source=str(path),
                purpose=purpose,
                file_type="docx",
                extracted_text=extracted_text,
                confidence=0.95 if extracted_text else 0.0,
                tables=table_rows,
                warnings=warnings,
            )
        except Exception as error:
            return DocumentExtractionResult(
                source=str(path),
                purpose=purpose,
                file_type="docx",
                confidence=0.0,
                errors=[f"docx_extraction_failed: {error}"],
            )
