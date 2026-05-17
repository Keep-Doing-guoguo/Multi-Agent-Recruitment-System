import os
from tempfile import TemporaryDirectory
import unittest
from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from recruitment_system.api import app


class ResumeApiTest(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = TemporaryDirectory()
        self.previous_db_path = os.environ.get("RECRUITMENT_DB_PATH")
        self.previous_enable_llm = os.environ.get("ENABLE_LLM")
        self.previous_llm_api_key = os.environ.get("LLM_API_KEY")
        os.environ["RECRUITMENT_DB_PATH"] = os.path.join(self.temp_dir.name, "test.sqlite3")
        os.environ["ENABLE_LLM"] = "false"
        self.client = TestClient(app)

    def tearDown(self) -> None:
        if self.previous_db_path is None:
            os.environ.pop("RECRUITMENT_DB_PATH", None)
        else:
            os.environ["RECRUITMENT_DB_PATH"] = self.previous_db_path
        if self.previous_enable_llm is None:
            os.environ.pop("ENABLE_LLM", None)
        else:
            os.environ["ENABLE_LLM"] = self.previous_enable_llm
        if self.previous_llm_api_key is None:
            os.environ.pop("LLM_API_KEY", None)
        else:
            os.environ["LLM_API_KEY"] = self.previous_llm_api_key
        self.temp_dir.cleanup()

    def create_conversation_with_resume_and_jd(
        self,
        resume_text: str = "张三\n本科，5 年 Python 后端开发经验\n技能: Python, FastAPI, Redis\n工作经历: 负责招聘系统 API",
        jd_text: str = "职位: Python 后端工程师\n要求: Python, FastAPI, Redis\n本科，3 年经验",
    ) -> dict:
        response = self.client.post(
            "/api/conversation/message",
            data={
                "message": "请分析这份简历是否匹配这个岗位",
                "jd_input": jd_text,
            },
            files={
                "resume_file": (
                    "resume.txt",
                    resume_text,
                    "text/plain",
                )
            },
        )
        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertTrue(body["success"])
        return body

    def test_parse_resume_upload_returns_candidate_profile(self) -> None:
        response = self.client.post(
            "/api/resume/parse",
            files={
                "file": (
                    "resume.txt",
                    "赵六\n本科，5 年 Python 后端开发经验\n技能: Python, FastAPI, PostgreSQL\n工作经历:\n负责 API 开发",
                    "text/plain",
                )
            },
        )

        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertTrue(body["success"])
        self.assertEqual(body["message"], "简历解析成功")
        self.assertEqual(body["data"]["candidate_profile"]["name"], "赵六")
        self.assertIn("Python", body["data"]["candidate_profile"]["skills"])

    def test_parse_resume_upload_returns_message_for_non_resume(self) -> None:
        response = self.client.post(
            "/api/resume/parse",
            files={
                "file": (
                    "note.txt",
                    "今天下午三点开会，讨论办公室采购和团建预算。",
                    "text/plain",
                )
            },
        )

        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertFalse(body["success"])
        self.assertIn("不像一份简历", body["message"])

    def test_parse_resume_upload_rejects_empty_file(self) -> None:
        response = self.client.post(
            "/api/resume/parse",
            files={"file": ("empty.txt", "", "text/plain")},
        )

        self.assertEqual(response.status_code, 400)

    def test_parse_resume_upload_accepts_docx_file(self) -> None:
        doc = Document()
        doc.add_paragraph("钱七")
        doc.add_paragraph("本科，4 年 Java 和 Spring Boot 后端开发经验")
        doc.add_paragraph("技能: Java, Spring Boot, MySQL, Redis")
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = self.client.post(
            "/api/resume/parse",
            files={
                "file": (
                    "resume.docx",
                    buffer.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertTrue(body["success"])
        self.assertEqual(body["data"]["document"]["file_type"], "docx")
        self.assertEqual(body["data"]["candidate_profile"]["name"], "钱七")

    def test_conversation_message_answers_from_existing_state(self) -> None:
        first_body = self.create_conversation_with_resume_and_jd(
            resume_text="李四\n本科，2 年 Python 后端开发经验\n技能: Python\n工作经历: 负责内部工具开发",
            jd_text="职位: Python 后端工程师\n要求: Python, FastAPI, Redis\n本科，5 年经验",
        )

        response = self.client.post(
            "/api/conversation/message/json",
            json={
                "conversation_id": first_body["conversation_id"],
                "message": "这个候选人为什么不推荐？",
            },
        )

        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertTrue(body["success"])
        self.assertEqual(body["route_decision"]["route"], "answer_from_state")
        self.assertTrue(body["message"])

    def test_conversation_message_requires_resume_file_for_first_business_run(self) -> None:
        response = self.client.post(
            "/api/conversation/message",
            data={
                "message": "我上传了一份简历，请解析",
            },
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("resume_file", response.json()["detail"])

    def test_json_message_requires_conversation_id(self) -> None:
        response = self.client.post(
            "/api/conversation/message/json",
            json={"message": "请分析这份简历是否匹配这个岗位"},
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("conversation_id", response.json()["detail"])

    def test_conversation_message_reports_missing_llm_key_when_llm_enabled(self) -> None:
        os.environ["ENABLE_LLM"] = "true"
        os.environ["LLM_API_KEY"] = ""

        response = self.client.post(
            "/api/conversation/message",
            data={
                "message": "请分析这份简历是否匹配这个岗位",
                "jd_input": "职位: Python 后端工程师\n要求: Python",
            },
            files={
                "resume_file": (
                    "resume.txt",
                    "吴十\n本科，7 年 Python 后端开发经验\n技能: Python",
                    "text/plain",
                )
            },
        )

        self.assertEqual(response.status_code, 500)
        self.assertIn("LLM_API_KEY", response.json()["detail"])

    def test_conversation_message_routes_first_resume_and_jd_to_full_flow(self) -> None:
        response = self.client.post(
            "/api/conversation/message",
            data={
                "message": "请分析这份简历是否匹配这个岗位",
                "jd_input": "职位: Python 后端工程师\n要求: Python, FastAPI, Redis\n本科，3 年经验",
            },
            files={
                "resume_file": (
                    "resume.txt",
                    "吴十\n本科，7 年 Python 后端开发经验\n技能: Python, FastAPI, Redis\n工作经历: 负责招聘系统 API",
                    "text/plain",
                )
            },
        )

        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertTrue(body["success"])
        self.assertEqual(body["route_decision"]["route"], "resume_intake")
        self.assertEqual(
            body["route_decision"]["required_nodes"],
            ["resume_intake", "resume_parsing", "jd_extraction", "job_matching", "screening", "interview", "supervisor"],
        )
        self.assertIn("candidate_profile", body["conversation_state"])
        self.assertIn("match_result", body["conversation_state"])
        self.assertIn("screening_result", body["conversation_state"])
        self.assertIn("interview_plan", body["conversation_state"])
        self.assertIn("supervisor_review", body["conversation_state"])
        self.assertIn("conversation_id", body)
        self.assertIn("run_id", body)

    def test_conversation_message_routes_jd_change_to_matching_and_downstream_agents(self) -> None:
        first_body = self.create_conversation_with_resume_and_jd(
            resume_text="周九\n本科，5 年 Python 后端开发经验\n技能: Python, FastAPI, Redis\n工作经历: 负责招聘系统 API",
            jd_text="职位: Python 后端工程师\n要求: Python, FastAPI\n本科，3 年经验",
        )

        response = self.client.post(
            "/api/conversation/message/json",
            json={
                "conversation_id": first_body["conversation_id"],
                "message": "JD 改成需要 FastAPI 和 Redis，请重新匹配",
                "jd_input": "职位: Python 后端工程师\n要求: Python, FastAPI, Redis\n本科，3 年经验",
            },
        )

        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertTrue(body["success"])
        self.assertEqual(body["route_decision"]["route"], "job_matching")
        self.assertIn("match_result", body["data"])
        self.assertIn("screening_result", body["data"])
        self.assertIn("interview_plan", body["data"])
        self.assertIn("supervisor_review", body["data"])

    def test_conversation_message_restores_state_from_database(self) -> None:
        first_body = self.create_conversation_with_resume_and_jd(
            resume_text="王十二\n本科，5 年 Python 后端开发经验\n技能: Python, FastAPI, Redis\n工作经历: 负责招聘系统 API",
            jd_text="职位: Python 后端工程师\n要求: Python, FastAPI\n本科，3 年经验",
        )

        second_response = self.client.post(
            "/api/conversation/message/json",
            json={
                "conversation_id": first_body["conversation_id"],
                "message": "JD 改成需要 Redis，请重新匹配",
                "jd_input": "职位: Python 后端工程师\n要求: Python, Redis\n本科，3 年经验",
            },
        )

        self.assertEqual(second_response.status_code, 200)
        second_body = second_response.json()
        self.assertTrue(second_body["success"])
        self.assertEqual(second_body["conversation_id"], first_body["conversation_id"])
        self.assertEqual(second_body["route_decision"]["route"], "job_matching")
        self.assertIn("match_result", second_body["data"])
        self.assertEqual(second_body["conversation_state"]["candidate_profile"]["name"], "王十二")

    def test_conversation_message_accepts_resume_file_and_persists_messages(self) -> None:
        response = self.client.post(
            "/api/conversation/message",
            data={
                "message": "请分析这份简历是否匹配这个岗位",
                "jd_input": "职位: Python 后端工程师\n要求: Python, FastAPI, Redis\n本科，3 年经验",
            },
            files={
                "resume_file": (
                    "resume.txt",
                    "郑十一\n本科，5 年 Python 后端开发经验\n技能: Python, FastAPI, Redis\n工作经历: 负责招聘系统 API",
                    "text/plain",
                )
            },
        )

        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertTrue(body["success"])
        self.assertEqual(body["route_decision"]["route"], "resume_intake")
        self.assertIn("conversation_id", body)
        self.assertIn("run_id", body)
        self.assertNotIn("resume_input", body["conversation_state"])
        self.assertEqual(body["conversation_state"]["candidate_profile"]["name"], "郑十一")

        conversation_response = self.client.get(f"/api/conversations/{body['conversation_id']}")
        self.assertEqual(conversation_response.status_code, 200)
        conversation_body = conversation_response.json()
        self.assertTrue(conversation_body["success"])
        self.assertEqual(conversation_body["conversation"]["id"], body["conversation_id"])
        self.assertEqual(len(conversation_body["messages"]), 2)
        self.assertEqual(conversation_body["messages"][0]["role"], "user")
        self.assertEqual(conversation_body["messages"][1]["role"], "assistant")


if __name__ == "__main__":
    unittest.main()
